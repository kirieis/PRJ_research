import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def style_heading(p, text, level, font_name="Segoe UI"):
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = font_name
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(27, 54, 93) # Navy
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(41, 128, 185) # Blue
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(52, 73, 94) # Dark Slate

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "F4F6F7")
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(44, 62, 80)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Create output directories
os.makedirs(os.path.join("docs", "reports"), exist_ok=True)
os.makedirs("reports", exist_ok=True)

# ---------------------------------------------------------
# GENERATE VIETNAMESE DOCX REPORT
# ---------------------------------------------------------
doc_vi = docx.Document()

# Page Setup
for section in doc_vi.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title Header
p_title = doc_vi.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("BÁO CÁO PHÂN TÍCH CODE VÀ TASK HOÀN THÀNH OF DEVELOPER 2\n(DATABASE ARCHITECTURE & FINANCIAL TRANSACTIONS)")
run_title.font.name = "Segoe UI"
run_title.font.size = Pt(18)
run_title.bold = True
run_title.font.color.rgb = RGBColor(27, 54, 93)

p_sub = doc_vi.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Dự án LUCY (Language Unity & Collaborative Youth) | Tác giả: Developer 2 | Cập nhật: Tháng 7/2026")
run_sub.font.name = "Segoe UI"
run_sub.font.size = Pt(10)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(127, 140, 141)

doc_vi.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 1
style_heading(doc_vi.add_paragraph(), "1. Tổng Quan & Bảng Công Việc Đã Hoàn Thành (Plan vs Implementation)", 1)

p_desc = doc_vi.add_paragraph()
p_desc.paragraph_format.space_after = Pt(6)
r = p_desc.add_run("Theo các tài liệu kế hoạch (plan.docx, alter_plan.docx, plan_after.docx), Developer 2 ban đầu phụ trách Cơ sở dữ liệu MS SQL Server. Do hoàn thành DB sớm hơn dự kiến, Dev 2 đã nhận thêm khối nhiệm vụ C# .NET Backend về Quản lý Tài chính (Wallet & Ledger), Xác thực Phòng ẩn danh (Anonymous Room Flow & Dual-Token JWT), Xử lý Bất đồng bộ / Race Condition & Deadlock, Audit Logging và Concurrency Testing.")
r.font.name = "Segoe UI"
r.font.size = Pt(10.5)

# Table 1: Completed Tasks Matrix
tbl_vi = doc_vi.add_table(rows=1, cols=4)
tbl_vi.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_vi.autofit = False

hdr_cells = tbl_vi.rows[0].cells
headers = ["Hạng Mục Task", "Chi Tiết Nhiệm Vụ Theo Kế Hoạch", "Kết Quả Thực Hiện / Deliverables", "Trạng Thái"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_shading(hdr_cells[i], "1B365D")
    p = hdr_cells[i].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.name = "Segoe UI"
    p.runs[0].font.size = Pt(9.5)
    set_cell_margins(hdr_cells[i])

tasks_vi = [
    ("Thiết Kế CSDL Cốt Lõi (Tuần 1-5)", "Thiết lập DB SQL Server LucyDB, map các bảng phân cấp bài học 6 phần (Languages -> Stages -> Levels -> SubLevels), tạo bảng QuizQuestions, AISupportQuestions, Rooms, Podcasts.", "Hoàn thành script LucyDB.sql và 002_week10_completion_schema.sql với ràng buộc FK và hệ thống Index tối ưu.", "Đã Hoàn Thành"),
    ("Tối Ưu AI Moderator (Tuần 6-7)", "Tạo Stored Procedure quét nhanh bộ câu hỏi gợi ý AI theo mốc thời gian (trigger_minute) để hỗ trợ phòng học realtime.", "Tạo Stored Procedure sp_GetAISupportByMinute chạy trực tiếp dưới DB, loại bỏ N+1 query.", "Đã Hoàn Thành"),
    ("Phòng Học Ẩn Danh & Token Kép (Tuần 1-5 Bổ Sung)", "Chuẩn hóa model Users, xây dựng bảng UserPersona (avatar, displayName ảo) và cấp RealtimeToken cho Node.js server.", "Hoàn thành AnonymousRoomAccessService.cs, RandomPersonaGenerator.cs, và JwtTokenService.cs với cơ chế Token kép không chứa PII.", "Đã Hoàn Thành"),
    ("Hệ Thống Ví & Sổ Cái Bất Biến (Tuần 6-7 Bổ Sung)", "Thiết kế CSDL Module tài chính: bảng Wallets (Ví), Transactions (Giao dịch), WalletLedger (Sổ cái lưu vết). Quy ước idempotency key.", "Tạo các script 003_wallets.sql, 004_wallet_transactions.sql, 005_wallet_ledger.sql với ràng buộc CHECK (balance >= 0) và idempotency key.", "Đã Hoàn Thành"),
    ("API Tặng Quà & Khóa Bi Quan (Tuần 8-9 Bổ Sung)", "Phát triển API POST /api/wallet/gift, chống race condition khi tặng quà đồng thời, ghi sổ cái trước/sau giao dịch.", "Hoàn thành WalletService.cs & SqlWalletRepository.cs sử dụng UPDLOCK, HOLDLOCK, ROWLOCK, sắp xếp ID ví tăng dần tránh deadlock, ghi ledger kép DEBIT/CREDIT.", "Đã Hoàn Thành"),
    ("Audit Logging & Concurrency Test (Tuần 10 Bổ Sung)", "Kiểm toán hệ thống ngoại lệ/logging và thực hiện script test đồng thời chuyên sâu cho gift và deposit.", "Tạo 006_audit_logs.sql, SqlAuditLogRepository.cs, script wallet-concurrency.ps1 (100 workers) và tài liệu docs/week10-concurrency-testing.md.", "Đã Hoàn Thành")
]

for item in tasks_vi:
    row_cells = tbl_vi.add_row().cells
    for i, text in enumerate(item):
        row_cells[i].text = text
        set_cell_margins(row_cells[i])
        p = row_cells[i].paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].font.name = "Segoe UI"
            p.runs[0].font.size = Pt(9)
            if i == 3:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(39, 174, 96)

# Section 2
style_heading(doc_vi.add_paragraph(), "2. Phân Tích Code Trong Project & Tác Dụng Của Mã Nguồn Dev 2", 1)

p_sec2 = doc_vi.add_paragraph()
p_sec2.paragraph_format.space_after = Pt(4)
r = p_sec2.add_run("Mã nguồn do Dev 2 trực tiếp xây dựng nằm ở 3 khu vực chính: Các Script CSDL SQL Server (data/SQL_database/ & lucy-auth-service/database/), Dịch vụ Backend C# ASP.NET Core (.NET Auth & Wallet Service), và Bộ kịch bản Concurrency Test.")
r.font.name = "Segoe UI"
r.font.size = Pt(10.5)

style_heading(doc_vi.add_paragraph(), "2.1 Khối Script Cơ Sở Dữ Liệu SQL Server", 2)
sql_files_vi = [
    ("data/SQL_database/LucyDB.sql & 002_week10_completion_schema.sql", "Định nghĩa toàn bộ cấu trúc CSDL bài học LUCY (languages, stages, levels, sub_levels, content_items, quiz_questions, rooms, podcasts, room_resources, ai_support_questions). Chứa các chỉ mục tối ưu latency (IX_levels_lang_stage, IX_sublevels_level, IX_contentitems_sub, IX_room_resources_active) và Stored Procedure sp_GetAISupportByMinute lấy câu hỏi AI theo thời gian thực."),
    ("lucy-auth-service/database/001_users_auth_contract.sql", "Tạo bảng dbo.users quản lý tài khoản người dùng, vai trò (LUCY, MENTOR, SUPER, ADMIN), password hash BCrypt, cờ is_anonymous."),
    ("lucy-auth-service/database/002_persona_realtime_contract.sql", "Tạo bảng dbo.user_personas lưu trữ danh tính ảo (display_name ngẫu nhiên, avatar_url, persona_type) cho luồng học thoại ẩn danh."),
    ("lucy-auth-service/database/003_wallets.sql", "Tạo bảng dbo.wallets quản lý ví tiền tệ của người dùng. Có ràng buộc CHECK (balance >= 0) chống số dư âm và trạng thái is_locked."),
    ("lucy-auth-service/database/004_wallet_transactions.sql", "Tạo bảng dbo.wallet_transactions lưu vết lịch sử giao dịch (DEPOSIT, GIFT, WITHDRAW) với khóa idempotency_key (128 ký tự) chống trùng lặp."),
    ("lucy-auth-service/database/005_wallet_ledger.sql", "Tạo bảng dbo.wallet_ledger đóng vai trò Sổ cái bất biến (Immutable Ledger). Mỗi giao dịch phát sinh 2 dòng (DEBIT cho ví gửi, CREDIT cho ví nhận) ghi nhận balance_before và balance_after."),
    ("lucy-auth-service/database/006_audit_logs.sql", "Tạo bảng dbo.audit_logs ghi nhật ký kiểm toán cho các thao tác bảo mật nhạy cảm và các ngoại lệ Unhandled Exception.")
]
for path, desc in sql_files_vi:
    p = doc_vi.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{path}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

style_heading(doc_vi.add_paragraph(), "2.2 Khối Mã Nguồn Backend C# ASP.NET Core (lucy-auth-service)", 2)
cs_files_vi = [
    ("Services/AnonymousRoomAccessService.cs & RandomPersonaGenerator.cs", "Xử lý nghiệp vụ truy cập phòng học ẩn danh. Kiểm tra role 'LUCY' để bắt buộc chế độ ẩn danh, gọi RandomPersonaGenerator sinh tên ảo (vd: 'Anonymous Fox') và lưu vào user_personas."),
    ("Services/JwtTokenService.cs & RsaJwtKeyProvider.cs", "Triển khai chiến lược Token Kép: Sinh AccessToken cho Mobile App và RealtimeToken (JWT ký bằng RSA/JWKS) dành riêng cho Node.js Socket server. RealtimeToken chỉ chứa thông tin ẩn danh (anon: true, display_name ảo, subject guid), không chứa PII (email, real user ID)."),
    ("Services/Wallet/WalletService.cs", "Trái tim của module tài chính. Chứa toàn bộ logic xử lý Nạp tiền (DepositAsync), Tặng quà (GiftAsync), và Truy vấn số dư (GetBalanceAsync). Đảm bảo tính Idempotency, khóa ví theo thứ tự ID tăng dần chống deadlock, và ghi sổ cái trước khi cập nhật số dư."),
    ("Data/Wallet/SqlWalletRepository.cs, SqlWalletTransactionRepository.cs, SqlWalletLedgerRepository.cs", "Tầng truy xuất dữ liệu ADO.NET thuần hiệu năng cao. Thực hiện các truy vấn SQL chứa khóa bi quan nghiêm ngặt (UPDLOCK, HOLDLOCK, ROWLOCK) để ngăn chặn race-condition."),
    ("Data/SqlAuditLogRepository.cs", "Thực thi ghi log kiểm toán sự kiện hệ thống và lỗi exception vào bảng dbo.audit_logs.")
]
for path, desc in cs_files_vi:
    p = doc_vi.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{path}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

style_heading(doc_vi.add_paragraph(), "2.3 Khối Kiểm Thử Đồng Thời (Concurrency Testing)", 2)
test_files_vi = [
    ("lucy-auth-service/tests/wallet-concurrency.ps1", "Kịch bản PowerShell tự động hóa kiểm thử tải đồng thời. Giả lập 100 luồng (Gift Workers & Deposit Workers) gửi request cùng lúc tới API để kiểm tra tính Idempotency, số dư ví không âm và không bị deadlock."),
    ("docs/week10-concurrency-testing.md", "Tài liệu quy định điều kiện tiên quyết, câu lệnh thực thi test và tiêu chuẩn nghiệm thu pass/fail hệ thống tài chính.")
]
for path, desc in test_files_vi:
    p = doc_vi.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{path}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

# Section 3
style_heading(doc_vi.add_paragraph(), "3. Phân Tích Chuyên Sâu Các Giải Pháp Kỹ Thuật Của Dev 2", 1)

style_heading(doc_vi.add_paragraph(), "3.1 Cơ Chế Chống Race-Condition & Tránh Deadlock Trong WalletService", 2)
p_exp1 = doc_vi.add_paragraph()
p_exp1.paragraph_format.space_after = Pt(4)
r = p_exp1.add_run("Để giải quyết bài toán giao dịch tài chính khi hàng trăm người dùng bấm 'Tặng quà' cùng một thời điểm, Dev 2 đã áp dụng 4 nguyên tắc kỹ thuật cốt lõi:")
r.font.name = "Segoe UI"
r.font.size = Pt(10)

principles_vi = [
    ("1. Idempotency Check First: ", "Hệ thống tra cứu idempotency_key với gợi ý khóa hàng UPDLOCK, HOLDLOCK. Nếu request đã được xử lý thành công trước đó, hệ thống lập tức trả về kết quả cũ mà không chạy lại logic giao dịch."),
    ("2. Lock Ordering (Sắp xếp thứ tự khóa tránh Deadlock): ", "Khi chuyển tiền giữa Ví A và Ví B, hệ thống luôn sắp xếp ID ví theo thứ tự tăng dần (math.Min / math.Max) trước khi xin khóa. Điều này triệt tiêu hoàn toàn rủi ro Deadlock vòng lặp khóa (Circular Deadlock)."),
    ("3. Pessimistic Locking (Khóa bi quan dưới DB): ", "Sử dụng truy vấn SQL trực tiếp với gợi ý khóa ROWLOCK, UPDLOCK, HOLDLOCK để độc quyền dòng dữ liệu ví đến khi Transaction kết thúc:"),
    ("4. Ledger-First Atomic Rollback: ", "Ghi cả 2 đầu dòng sổ cái (DEBIT và CREDIT) vào dbo.wallet_ledger trước, sau đó mới cập nhật số dư bảng dbo.wallets. Bất kỳ lỗi nào phát sinh sẽ kích hoạt Rollback toàn bộ transaction.")
]
for title, desc in principles_vi:
    p = doc_vi.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

add_code_block(doc_vi, "SELECT id, user_id, balance, currency, is_locked, created_at, updated_at\nFROM dbo.wallets WITH (UPDLOCK, HOLDLOCK, ROWLOCK)\nWHERE id = @walletId;")

style_heading(doc_vi.add_paragraph(), "3.2 Stored Procedure Tối Ưu AI Moderator (sp_GetAISupportByMinute)", 2)
add_code_block(doc_vi, "CREATE PROCEDURE sp_GetAISupportByMinute\n    @sub_level_id   INT,\n    @current_minute INT\nAS\nBEGIN\n    SET NOCOUNT ON;\n    SELECT TOP 3 id, sub_level_id, question_text, trigger_minute, language_id, order_index\n    FROM ai_support_questions\n    WHERE sub_level_id   = @sub_level_id\n      AND trigger_minute <= @current_minute\n    ORDER BY trigger_minute DESC, order_index ASC;\nEND;")

# Section 4
style_heading(doc_vi.add_paragraph(), "4. Kết Quả Kiểm Thử Đồng Thời & Kết Luận", 1)
p_sec4 = doc_vi.add_paragraph()
p_sec4.paragraph_format.space_after = Pt(4)
r = p_sec4.add_run("Kết quả kiểm thử đồng thời (Concurrency Test) thông qua script wallet-concurrency.ps1 xác nhận các tiêu chí nghiệm thu đều đạt 100%:")
r.font.name = "Segoe UI"
r.font.size = Pt(10)

pass_criteria_vi = [
    "Không phát sinh số dư âm trên bất kỳ ví nào (CHECK balance >= 0 được bảo vệ tuyệt đối).",
    "Khóa Idempotency_key duy nhất 100%, không bị lặp giao dịch dù nhận 100 request trùng key.",
    "Mỗi giao dịch Gift thành công có đúng 1 bản ghi DEBIT và 1 bản ghi CREDIT trên sổ cái wallet_ledger.",
    "Không xảy ra lỗi Deadlock dưới SQL Server nhờ cơ chế Lock Ordering theo ID ví tăng dần.",
    "Toàn bộ exception được ghi vết đầy đủ vào bảng audit_logs mà không làm lệch số dư."
]
for item in pass_criteria_vi:
    p = doc_vi.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(item)
    r.font.name = "Segoe UI"
    r.font.size = Pt(9.5)

p_conclusion_vi = doc_vi.add_paragraph()
p_conclusion_vi.paragraph_format.space_before = Pt(8)
r_conc = p_conclusion_vi.add_run("KẾT LUẬN: Developer 2 đã hoàn thành 100% khối lượng công việc theo cả kế hoạch gốc và kế hoạch điều chỉnh, xây dựng thành công nền tảng Cơ sở dữ liệu vững chắc và Module Tài chính / Phòng ẩn danh đạt chuẩn an toàn cao nhất trong dự án LUCY.")
r_conc.bold = True
r_conc.font.name = "Segoe UI"
r_conc.font.size = Pt(10)
r_conc.font.color.rgb = RGBColor(27, 54, 93)

# Save Vietnamese DOCX
doc_vi.save(os.path.join("docs", "reports", "dev2_report.docx"))
doc_vi.save(os.path.join("reports", "dev2_report.docx"))

print("Generated dev2_report.docx successfully!")


# ---------------------------------------------------------
# GENERATE ENGLISH DOCX REPORT
# ---------------------------------------------------------
doc_en = docx.Document()

# Page Setup
for section in doc_en.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title Header
p_title_en = doc_en.add_paragraph()
p_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title_en = p_title_en.add_run("CODE ANALYSIS & TASK COMPLETION REPORT: DEVELOPER 2\n(DATABASE ARCHITECTURE & FINANCIAL TRANSACTIONS)")
run_title_en.font.name = "Segoe UI"
run_title_en.font.size = Pt(18)
run_title_en.bold = True
run_title_en.font.color.rgb = RGBColor(27, 54, 93)

p_sub_en = doc_en.add_paragraph()
p_sub_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub_en = p_sub_en.add_run("LUCY Project (Language Unity & Collaborative Youth) | Author: Developer 2 | Date: July 2026")
run_sub_en.font.name = "Segoe UI"
run_sub_en.font.size = Pt(10)
run_sub_en.font.italic = True
run_sub_en.font.color.rgb = RGBColor(127, 140, 141)

doc_en.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 1
style_heading(doc_en.add_paragraph(), "1. Overview & Completed Tasks Matrix (Plan vs Implementation)", 1)

p_desc_en = doc_en.add_paragraph()
p_desc_en.paragraph_format.space_after = Pt(6)
r = p_desc_en.add_run("According to the project planning documents (plan.docx, alter_plan.docx, plan_after.docx), Developer 2 originally focused on MS SQL Server Database development. Having completed the database tasks ahead of schedule, Developer 2 was allocated additional backend duties in the C# .NET service, covering Financial Transactions (Wallet & Ledger), Anonymous Room Access & Dual-Token Strategy, Concurrency & Deadlock Control, Audit Logging, and Concurrency Testing.")
r.font.name = "Segoe UI"
r.font.size = Pt(10.5)

# Table 1: Completed Tasks Matrix (English)
tbl_en = doc_en.add_table(rows=1, cols=4)
tbl_en.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_en.autofit = False

hdr_cells_en = tbl_en.rows[0].cells
headers_en = ["Task Category", "Planned Task Details", "Deliverables & Implementation", "Status"]
for i, h in enumerate(headers_en):
    hdr_cells_en[i].text = h
    set_cell_shading(hdr_cells_en[i], "1B365D")
    p = hdr_cells_en[i].paragraphs[0]
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.name = "Segoe UI"
    p.runs[0].font.size = Pt(9.5)
    set_cell_margins(hdr_cells_en[i])

tasks_en = [
    ("Core DB Architecture (Weeks 1-5)", "Set up SQL Server LucyDB, map 6-level lesson hierarchy (Languages -> Stages -> Levels -> SubLevels), create QuizQuestions, AISupportQuestions, Rooms, Podcasts.", "Completed LucyDB.sql and 002_week10_completion_schema.sql with FK constraints and optimized indexes.", "Completed"),
    ("AI Moderator Optimization (Weeks 6-7)", "Create Stored Procedure to scan AI support questions by time (trigger_minute) for real-time rooms.", "Implemented sp_GetAISupportByMinute SP running directly in DB to eliminate N+1 queries.", "Completed"),
    ("Anonymous Room Flow & Dual Token (Weeks 1-5 Adj.)", "Standardize Users model, build UserPersona table (random displayName, avatar) and issue RealtimeToken for Node.js server.", "Completed AnonymousRoomAccessService.cs, RandomPersonaGenerator.cs, and JwtTokenService.cs with PII-free RSA/JWKS Dual Token architecture.", "Completed"),
    ("Wallet & Ledger Schema (Weeks 6-7 Adj.)", "Design financial CSDL: Wallets, Transactions, WalletLedger tables + idempotency key conventions.", "Created 003_wallets.sql, 004_wallet_transactions.sql, 005_wallet_ledger.sql with CHECK (balance >= 0) and unique idempotency constraints.", "Completed"),
    ("Gift API & Concurrency Control (Weeks 8-9 Adj.)", "Develop POST /api/wallet/gift API, prevent race conditions during concurrent gifting, write double-entry ledger.", "Implemented WalletService.cs & SqlWalletRepository.cs with UPDLOCK, HOLDLOCK, ROWLOCK, stable lock ordering, and atomic DEBIT/CREDIT ledger writes.", "Completed"),
    ("Audit Logging & Concurrency Testing (Week 10 Adj.)", "Audit exception/logging system and execute concurrency test scripts for gift and deposit flows.", "Created 006_audit_logs.sql, SqlAuditLogRepository.cs, wallet-concurrency.ps1 (100 workers), and docs/week10-concurrency-testing.md.", "Completed")
]

for item in tasks_en:
    row_cells = tbl_en.add_row().cells
    for i, text in enumerate(item):
        row_cells[i].text = text
        set_cell_margins(row_cells[i])
        p = row_cells[i].paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].font.name = "Segoe UI"
            p.runs[0].font.size = Pt(9)
            if i == 3:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(39, 174, 96)

# Section 2 (English)
style_heading(doc_en.add_paragraph(), "2. Codebase Mapping & Component Purpose", 1)

p_sec2_en = doc_en.add_paragraph()
p_sec2_en.paragraph_format.space_after = Pt(4)
r = p_sec2_en.add_run("The codebase developed by Developer 2 spans 3 primary modules: SQL Server Database Scripts (data/SQL_database/ & lucy-auth-service/database/), Backend C# ASP.NET Core (.NET Auth & Wallet Service), and Concurrency Test Suites.")
r.font.name = "Segoe UI"
r.font.size = Pt(10.5)

style_heading(doc_en.add_paragraph(), "2.1 SQL Server Database Scripts", 2)
sql_files_en = [
    ("data/SQL_database/LucyDB.sql & 002_week10_completion_schema.sql", "Defines full LUCY lesson hierarchy (languages, stages, levels, sub_levels, content_items, quiz_questions, rooms, podcasts, room_resources, ai_support_questions). Contains latency-optimized indexes (IX_levels_lang_stage, IX_sublevels_level, IX_contentitems_sub, IX_room_resources_active) and sp_GetAISupportByMinute SP."),
    ("lucy-auth-service/database/001_users_auth_contract.sql", "Creates dbo.users table for user accounts, role check constraints (LUCY, MENTOR, SUPER, ADMIN), BCrypt password hashes, and is_anonymous flags."),
    ("lucy-auth-service/database/002_persona_realtime_contract.sql", "Creates dbo.user_personas table storing virtual identities (random display_name, avatar_url, persona_type) for anonymous audio rooms."),
    ("lucy-auth-service/database/003_wallets.sql", "Creates dbo.wallets table with CHECK (balance >= 0) constraints and is_locked flag to manage user currency."),
    ("lucy-auth-service/database/004_wallet_transactions.sql", "Creates dbo.wallet_transactions table recording transaction history (DEPOSIT, GIFT, WITHDRAW) with 128-char idempotency_key."),
    ("lucy-auth-service/database/005_wallet_ledger.sql", "Creates dbo.wallet_ledger as an Immutable Ledger. Each transaction generates 2 rows (DEBIT for sender, CREDIT for recipient) recording balance_before and balance_after."),
    ("lucy-auth-service/database/006_audit_logs.sql", "Creates dbo.audit_logs table to audit sensitive security operations and unhandled application exceptions.")
]
for path, desc in sql_files_en:
    p = doc_en.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{path}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

style_heading(doc_en.add_paragraph(), "2.2 Backend C# ASP.NET Core Services (lucy-auth-service)", 2)
cs_files_en = [
    ("Services/AnonymousRoomAccessService.cs & RandomPersonaGenerator.cs", "Handles anonymous audio room access logic. Enforces anonymity for 'LUCY' role users, invoking RandomPersonaGenerator to assign virtual names (e.g., 'Anonymous Fox') into user_personas."),
    ("Services/JwtTokenService.cs & RsaJwtKeyProvider.cs", "Implements Dual-Token Security: Generates AccessToken for Mobile App and RealtimeToken (RSA/JWKS-signed JWT) for Node.js Socket server. RealtimeToken contains only anonymous claims (anon: true, fake display_name, fake subject guid) without PII."),
    ("Services/Wallet/WalletService.cs", "Core financial engine containing DepositAsync, GiftAsync, and GetBalanceAsync logic. Enforces idempotency checks, ascending wallet ID locking order against deadlocks, and double-entry ledger writes prior to balance mutation."),
    ("Data/Wallet/SqlWalletRepository.cs, SqlWalletTransactionRepository.cs, SqlWalletLedgerRepository.cs", "High-performance ADO.NET data access layer executing raw SQL with strict locking hints (UPDLOCK, HOLDLOCK, ROWLOCK) to eliminate race conditions."),
    ("Data/SqlAuditLogRepository.cs", "Executes audit logging for system security events and unhandled exceptions into dbo.audit_logs.")
]
for path, desc in cs_files_en:
    p = doc_en.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{path}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

style_heading(doc_en.add_paragraph(), "2.3 Concurrency Testing Suite", 2)
test_files_en = [
    ("lucy-auth-service/tests/wallet-concurrency.ps1", "PowerShell automated load testing script simulating 100 concurrent workers (Gift & Deposit Workers) calling APIs simultaneously to verify idempotency, non-negative balances, and deadlock freedom."),
    ("docs/week10-concurrency-testing.md", "Specification document outlining preconditions, test commands, and SQL verification pass/fail criteria for financial testing.")
]
for path, desc in test_files_en:
    p = doc_en.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{path}: ")
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

# Section 3 (English)
style_heading(doc_en.add_paragraph(), "3. Technical Deep-Dive into Developer 2 Solutions", 1)

style_heading(doc_en.add_paragraph(), "3.1 Race Condition & Deadlock Prevention in WalletService", 2)
p_exp1_en = doc_en.add_paragraph()
p_exp1_en.paragraph_format.space_after = Pt(4)
r = p_exp1_en.add_run("To ensure financial data integrity when hundreds of users simultaneously send gifts, Developer 2 established 4 technical safeguards:")
r.font.name = "Segoe UI"
r.font.size = Pt(10)

principles_en = [
    ("1. Idempotency Check First: ", "Queries idempotency_key with UPDLOCK, HOLDLOCK hints. If the transaction was completed previously, returns cached result instantly without executing financial logic."),
    ("2. Lock Ordering (Deadlock Elimination): ", "When transferring funds between Wallet A and Wallet B, wallet IDs are always locked in ascending order (Math.Min / Math.Max). This completely eliminates circular deadlock risks."),
    ("3. Pessimistic Row Locking: ", "Executes raw SQL queries with ROWLOCK, UPDLOCK, HOLDLOCK hints to acquire exclusive row ownership until transaction commit:"),
    ("4. Ledger-First Atomic Rollback: ", "Writes both DEBIT and CREDIT rows to dbo.wallet_ledger first, then mutates balances in dbo.wallets. Any failure triggers an immediate complete transaction rollback.")
]
for title, desc in principles_en:
    p = doc_en.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.name = "Segoe UI"
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(9.5)

add_code_block(doc_en, "SELECT id, user_id, balance, currency, is_locked, created_at, updated_at\nFROM dbo.wallets WITH (UPDLOCK, HOLDLOCK, ROWLOCK)\nWHERE id = @walletId;")

style_heading(doc_en.add_paragraph(), "3.2 AI Moderator Optimization Stored Procedure (sp_GetAISupportByMinute)", 2)
add_code_block(doc_en, "CREATE PROCEDURE sp_GetAISupportByMinute\n    @sub_level_id   INT,\n    @current_minute INT\nAS\nBEGIN\n    SET NOCOUNT ON;\n    SELECT TOP 3 id, sub_level_id, question_text, trigger_minute, language_id, order_index\n    FROM ai_support_questions\n    WHERE sub_level_id   = @sub_level_id\n      AND trigger_minute <= @current_minute\n    ORDER BY trigger_minute DESC, order_index ASC;\nEND;")

# Section 4 (English)
style_heading(doc_en.add_paragraph(), "4. Concurrency Test Verification & Conclusion", 1)
p_sec4_en = doc_en.add_paragraph()
p_sec4_en.paragraph_format.space_after = Pt(4)
r = p_sec4_en.add_run("Automated concurrency test results via wallet-concurrency.ps1 confirmed 100% compliance across all acceptance criteria:")
r.font.name = "Segoe UI"
r.font.size = Pt(10)

pass_criteria_en = [
    "Zero negative wallet balances across all scenarios (CHECK balance >= 0 strictly enforced).",
    "100% unique Idempotency_key enforcement, preventing duplicate execution under 100 parallel requests.",
    "Every successful gift transaction created exactly 1 DEBIT and 1 CREDIT entry in wallet_ledger.",
    "Zero SQL Server deadlocks encountered due to ascending wallet ID lock ordering.",
    "All exceptions cleanly captured in audit_logs without leaving orphaned ledger rows or corrupted balances."
]
for item in pass_criteria_en:
    p = doc_en.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(item)
    r.font.name = "Segoe UI"
    r.font.size = Pt(9.5)

p_conclusion_en = doc_en.add_paragraph()
p_conclusion_en.paragraph_format.space_before = Pt(8)
r_conc_en = p_conclusion_en.add_run("CONCLUSION: Developer 2 has successfully completed 100% of planned and adjusted tasks, establishing a robust database foundation, anonymous authentication pipeline, and bank-grade financial transaction module for the LUCY platform.")
r_conc_en.bold = True
r_conc_en.font.name = "Segoe UI"
r_conc_en.font.size = Pt(10)
r_conc_en.font.color.rgb = RGBColor(27, 54, 93)

# Save English DOCX
doc_en.save(os.path.join("docs", "reports", "dev2_report_en.docx"))
doc_en.save(os.path.join("reports", "dev2_report_en.docx"))

print("Generated dev2_report_en.docx successfully!")
